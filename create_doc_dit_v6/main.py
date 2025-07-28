import subprocess
import sys
import os
from pathlib import Path
from datetime import datetime
import logging
import re
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI

# Carrega variáveis do .env
load_dotenv()

llm = ChatOpenAI(
    model="gpt-4",
    temperature=0,
    openai_api_key=os.getenv("OPENAI_API_KEY")
)

markdown_path = Path("./markdown")
output_path = Path("./doc/documento_gerado.docx")

def extrair_blocos_codigo(md_content):
    return re.findall(r"```python(.*?)```", md_content, re.DOTALL)

def extrair_texto_sem_codigo(md_content):
    # Remove blocos de código para enviar só texto puro
    return re.sub(r"```.*?```", "", md_content, flags=re.DOTALL).strip()

def gerar_resumo_por_arquivo(titulo, codigo):
    prompt = f"""
Você é um assistente técnico.

Este é o código extraído do notebook "{titulo}":

{codigo}

Por favor, gere um parágrafo explicativo claro e objetivo sobre a lógica e o propósito do código acima.
"""
    resposta = llm.invoke(prompt)
    return resposta.content.strip()

def criar_doc_com_conteudo(md_files):
    doc = Document('./doc/DIT_model.docx')

    # Título principal
    doc.add_heading("Documentação Técnica Gerada", level=1)
    doc.add_paragraph(
        "Este documento foi gerado automaticamente a partir dos arquivos Markdown, "
        "seguindo o estilo e organização do modelo DIT_model.docx."
    )

    for md_file in md_files:
        titulo = md_file.stem.replace("_", " ").title()
        with open(md_file, "r", encoding="utf-8") as f:
            conteudo_md = f.read()

        # Extrai código e texto
        codigos = extrair_blocos_codigo(conteudo_md)
        texto_limpo = extrair_texto_sem_codigo(conteudo_md)

        # Adiciona título da seção
        doc.add_page_break()
        doc.add_heading(titulo, level=1)

        # Adiciona o texto explicativo / conteúdo textual
        if texto_limpo:
            doc.add_paragraph(texto_limpo)

        # Para cada bloco de código, gera resumo e adiciona no doc
        for codigo in codigos:
            resumo = gerar_resumo_por_arquivo(titulo, codigo)
            doc.add_heading("Resumo do Código", level=2)
            doc.add_paragraph(resumo)
            doc.add_heading("Código Fonte", level=2)

            # Adiciona bloco de código formatado
            p = doc.add_paragraph()
            run = p.add_run(codigo.strip())
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

    # Salva o documento final
    doc.save(output_path)
    print(f"✅ Documento criado em: {output_path}")
    
# ---------------------------------------------------------------------------------------------------------------------
def main():
    try:
        import nbconvert
    except ImportError:
        print("[ERRO] nbconvert não está instalado. Tentando instalar...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "nbconvert"])
        except Exception as install_error:
            print(f"[ERRO] Falha ao instalar nbconvert: {install_error}")
            sys.exit(1)
        
    from pathlib import Path
    import os
    base_dir = Path(__file__).resolve().parent
    sys.path.insert(0, str(base_dir / "functions"))

    from functions.log import configurar_logger
    from functions.estrutura import criar_pastas as criar_pastas_dinamico
    from functions.conversao import converte_to_md
    from functions.upsert_key_gpt import upsert_key_gpt
    from functions.create_key_gpt import create_key_gpt

    global logger
    logger = configurar_logger(base_dir)

    logger.info("[OK] Biblioteca 'nbconvert' disponível.")
    logger.info("[OK] Iniciando conversão de notebooks para Markdown...")


    notebooks_dir = base_dir / "notebooks"
    notebooks = [f for f in notebooks_dir.glob("*.ipynb") if f.is_file()]

    if not notebooks:
        logger.warning("[!] Nenhum arquivo .ipynb encontrado em /notebooks")
        return

    total = len(notebooks)
    convertidos = 0
    falhas = 0

    for arquivo in notebooks:
        sucesso = converte_to_md(arquivo, base_dir, logger)
        if sucesso:
            convertidos += 1
        else:
            falhas += 1

    logger.info("========== RESUMO DA EXECUÇÃO ========")
    logger.info(f"Total de arquivos encontrados______: {total}")
    logger.info(f"Convertidos com sucesso____________: {convertidos}")
    logger.info(f"Falharam na conversão______________: {falhas}")
    logger.info("======================================\n")

    print()
    resposta = input("Deseja criar um arquivo Dit? (S/n): ").strip().lower()
    if resposta in ["s", "sim", ""]:
        
        # Criação do arquivo de chave da OpenAI
        create_key_gpt(base_dir)
    
        dit_path = base_dir / "doc" / "notebooks.docx"
        with open(dit_path, "w", encoding="utf-8") as dit_file:
            for arquivo in notebooks:
                dit_file.write(f"{arquivo.name}\n")
        logger.info(f"[+] Arquivo Dit criado: {dit_path}")
    
    else:
        logger.info("[=] Criação do arquivo Dit cancelada.")

    # Creaçaõ do arquivo DIT final
    if not markdown_path.exists():
        print(f"Pasta {markdown_path} não encontrada!")
        return

    md_files = sorted(markdown_path.glob("*.md"))
    if not md_files:
        print(f"Nenhum arquivo Markdown (.md) encontrado em {markdown_path}")
        return

    criar_doc_com_conteudo(md_files)

# ---------------------------------------------------------------------------------------------------------------------
if __name__ == "__main__":
    main()
