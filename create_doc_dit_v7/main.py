import subprocess
import sys
import os
from pathlib import Path
import logging
import re
from docx import Document
from docx.oxml.ns import qn 
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from docx.shared import RGBColor

# Carrega variáveis do .env
#---------------------------------------------------------------------------------------------------------------------------------
load_dotenv()

# Inicializa LLM com chave
#---------------------------------------------------------------------------------------------------------------------------------
llm = ChatOpenAI(
    model="gpt-3.5-turbo-0125", # ou "gpt-4" se preferir
    temperature=0,
    openai_api_key=os.getenv("OPENAI_API_KEY")
)

# Caminhos
#---------------------------------------------------------------------------------------------------------------------------------
markdown_path = Path("./markdown")
output_path = Path("./doc/DIT.docx")
modelo_path = Path("./doc/DIT_model.docx")  # usado como referência

# Funções auxiliares
#---------------------------------------------------------------------------------------------------------------------------------
def extrair_blocos_codigo(md_content):
    return re.findall(r"```python(.*?)```", md_content, re.DOTALL)

#---------------------------------------------------------------------------------------------------------------------------------
def extrair_texto_sem_codigo(md_content):
    return re.sub(r"```.*?```", "", md_content, flags=re.DOTALL).strip()

#---------------------------------------------------------------------------------------------------------------------------------
from pathlib import Path

def gerar_sumario(md_files_list):
    md_nomes_formatados = "\n".join(
        Path(md).stem.replace("_", " ").title() for md in md_files_list
    )

    prompt = f"""
Você é um assistente técnico responsável por montar sumários técnicos.

Seu objetivo é gerar um **sumário estruturado** em formato de lista, baseado nos nomes dos arquivos Markdown fornecidos.

Caso os nomes dos arquivos não estejam em um formato adequado, você deve criar títulos apropriados, retirando números no início e ajustando capitalização.

Regras:
- Não repita o título "Sumário".
- Não repita o número duas vezes.
- Não inclua conteúdo extra nem explicações.
- Cada título deve estar em uma nova linha.
- Numerar os itens sequencialmente.

Aqui estão os nomes dos arquivos:
{md_nomes_formatados}

Gere apenas o sumário numerado nesse formato.
"""
    resposta = llm.invoke(prompt)
    return resposta.content.strip()


#---------------------------------------------------------------------------------------------------------------------------------
def gerar_introducao(md_files_list):
    # Transformar a lista em uma lista formatada para leitura
    arquivos_formatados = '\n'.join([f"- {nome}" for nome in md_files_list])

    prompt = f"""
Você é um assistente técnico especializado na elaboração de Documentos de Implementação Técnica (DIT). Seu objetivo é redigir a **introdução** do documento com base na lista de notebooks utilizados no projeto:

{arquivos_formatados}

### Regras de estilo:

- Linguagem: técnica, objetiva, institucional e clara.
- Título e subtítulos: já definidos; **NÃO os gere novamente**.
- Fonte do corpo do texto: "Arial Nova" (ou similar), tamanho 11, cor preta.
- Scripts Python: destaque com formatação de código, fonte "Arial Nova", tamanho 9, cor levemente acinzentada, espaçamento 1,15.
- Parágrafos: espaçamento entre linhas de 1,15, justificados.

### Objetivo:

Crie uma **introdução técnica padronizada**, que explique brevemente:
- O escopo geral do projeto.
- O objetivo do documento.
- O papel dos notebooks listados (por exemplo, se realizam ingestão, tratamento, análise ou visualização de dados).
- O valor ou impacto que esse documento oferece à equipe técnica ou ao negócio.

Evite repetições, **não escreva o título novamente**, e utilize linguagem clara, técnica e concisa.
"""
    resposta = llm.invoke(prompt)
    return resposta.content.strip()

#---------------------------------------------------------------------------------------------------------------------------------
def gerar_resumo_por_arquivo(titulo, codigo, index):
    prompt = f"""
Você é um assistente técnico responsável por gerar conteúdo para um Documento de Implementação Técnica (DIT), no padrão institucional da Minerva Foods.

Seu objetivo é produzir um parágrafo de resumo técnico para um trecho de código Python, seguindo a estrutura e estilo de documentação formal do modelo DIT.

### Instruções de formatação e estilo:

- O conteúdo faz parte da seção "{index}. {titulo}" do documento.
- Use linguagem **técnica**, **objetiva**, **profissional** e em **tom institucional**.
- O parágrafo deve ter:
  - Fonte: Arial Nova (ou similar), tamanho 11
  - Cor: preta
  - Espaçamento entre linhas: 1,15
- Os títulos e subtítulos do documento serão formatados com coloração azul tecnológica: #2D5BFF.
- O bloco de código será inserido com:
  - Fonte: Arial Nova (ou similar), tamanho 9
  - Cor levemente mais clara (cinza escuro)
  - Espaçamento entre linhas de 1,15
- Não repita o código, apenas gere o resumo.

### Trecho de código do notebook "{titulo}":

{codigo}

Com base nisso, escreva **apenas o parágrafo explicativo** sobre a lógica e propósito do código. Não inclua o código novamente.
"""
    resposta = llm.invoke(prompt)
    return resposta.content.strip()

#---------------------------------------------------------------------------------------------------------------------------------
def limpar_linhas_irrelevantes(texto):
    return "\n".join(
        linha for linha in texto.splitlines() if not linha.strip().startswith(("|", "#"))
    )

#---------------------------------------------------------------------------------------------------------------------------------
# Geração do documento estruturado com base no modelo
def criar_doc_com_conteudo(md_files):
    # novo doc em branco
    doc = Document() 
    
    # CAPA
    # Adiciona parágrafos vazios para empurrar o conteúdo para baixo (~30%)
    for _ in range(10):  # ajuste esse número conforme necessário
        doc.add_paragraph()
    
    # CAPA
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = titulo.add_run("Dataside")
    run.bold = True
    run.font.size = Pt(25)
    run.font.color.rgb = RGBColor(45, 91, 255)  # Azul tecnológico
    run.font.name = "Arial Nova"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Nova')  # Para garantir exibição correta em alguns ambientes

    # Subtítulo
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_sub = subtitulo.add_run("Documentação Técnica - Automação de Notebooks\n\nDocumento de Implementação Técnica")
    run_sub.font.size = Pt(14)
    run.font.name = "Arial Nova"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Nova')  # Para garantir exibição correta em alguns ambientes 

    doc.add_page_break()

    # INTRODUÇÃO
    doc.add_heading("Objetivo do documento", level=1)
    doc.add_paragraph(
        "Este Documento de Implementação Técnica (DIT) tem como objetivo registrar, de forma detalhada, padronizada e"
        "profissional, as soluções técnicas desenvolvidas no escopo do projeto, com ênfase na automação e execução de scripts"
        "contidos em notebooks."

        "O documento visa garantir rastreabilidade, compreensão técnica, reprodutibilidade das implementações e alinhamento com as" 
        "melhores práticas de desenvolvimento e documentação adotadas pela Dataside. Além disso, busca facilitar a comunicação"
        "entre equipes técnicas e não técnicas, assegurando que o conhecimento gerado esteja devidamente estruturado e acessível" 
        "para auditorias, manutenções futuras e expansão do projeto."
    )

    doc.add_page_break()
    
    # Adiciona sumário
    md_files_list = [f.name for f in md_files if f.is_file()]
    doc.add_heading("Sumário", level=1)
    sumario = gerar_sumario(md_files_list)

    for linha in sumario.splitlines():
        if linha.strip():  # evita adicionar linhas vazias
            doc.add_paragraph(linha, style='List Number')
    doc.add_page_break()
    
    # Adiciona Introdução
    doc.add_heading("Introdução", level=1)
    introducao = gerar_introducao(md_files_list)
    doc.add_paragraph(introducao)
    doc.add_page_break()

    # Adiciona conteúdo de cada markdown
    for idx, md_file in enumerate(md_files, start=1):
        titulo = md_file.stem.replace("_", "_").title()
        with open(md_file, "r", encoding="utf-8") as f:
            conteudo_md = limpar_linhas_irrelevantes(f.read())

        texto = extrair_texto_sem_codigo(conteudo_md)
        codigos = extrair_blocos_codigo(conteudo_md)

        doc.add_page_break()
        doc.add_heading(f"{idx}. {titulo}", level=1)

        if texto:
            doc.add_paragraph(texto)
            
        for i, codigo in enumerate(codigos, start=1):
            resumo = gerar_resumo_por_arquivo(titulo, codigo, f"{idx}.{i}")
            doc.add_heading(f"\n{idx}.{i} Resumo do Código", level=3)
            doc.add_paragraph(resumo)

            doc.add_heading("Código Fonte", level=4)
            p = doc.add_paragraph()
            run = p.add_run(codigo.strip())
            run.font.name = 'Arial Nova'
            run.font.size = Pt(9)


    # Rodapé institucional (não é rodapé técnico)
    doc.add_paragraph(
        "\n\nDocumento Interno - A divulgação sem autorização prévia viola as normas e diretrizes da organização.",
        style="Normal"
    )

    doc.save(output_path)
    print(f"✅ Documento criado com sucesso: {output_path}")

# Função principal
#---------------------------------------------------------------------------------------------------------------------------------
def main():
    try:
        import nbconvert
    except ImportError:
        print("[ERRO] nbconvert não está instalado. Instalando...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "nbconvert"])
        except Exception as install_error:
            print(f"[ERRO] Falha ao instalar nbconvert: {install_error}")
            sys.exit(1)

    base_dir = Path(__file__).resolve().parent
    sys.path.insert(0, str(base_dir / "functions"))

    from functions.log import configurar_logger
    from functions.estrutura import criar_pastas as criar_pastas_dinamico
    from functions.conversao import converte_to_md
    from functions.upsert_key_gpt import upsert_key_gpt
    from functions.create_key_gpt import create_key_gpt

    global logger
    logger = configurar_logger(base_dir)

    logger.info("[OK] Iniciando conversão de notebooks para Markdown...")

    notebooks_dir = base_dir / "notebooks"
    notebooks = [f for f in notebooks_dir.glob("*.ipynb") if f.is_file()]
    if not notebooks:
        logger.warning("Nenhum arquivo .ipynb encontrado.")
        return

    total, convertidos, falhas = len(notebooks), 0, 0
    for arquivo in notebooks:
        sucesso = converte_to_md(arquivo, base_dir, logger)
        if sucesso:
            convertidos += 1
        else:
            falhas += 1

    logger.info("========== RESUMO ==========")
    logger.info(f"Total: {total}, Sucesso: {convertidos}, Falhas: {falhas}")
    logger.info("============================")

    print()
    resposta = input("Deseja criar um arquivo DIT final? (S/n): ").strip().lower()
    if resposta in ["s", "sim", ""]:
        create_key_gpt(base_dir)
        dit_path = base_dir / "doc" / "notebooks.docx"
        with open(dit_path, "w", encoding="utf-8") as dit_file:
            for arquivo in notebooks:
                dit_file.write(f"{arquivo.name}\n")
        logger.info(f"[+] Arquivo Dit criado: {dit_path}")
    else:
        logger.info("[=] Criação do arquivo Dit cancelada.")

    # Agora criamos o documento final
    if not markdown_path.exists():
        print(f"Pasta {markdown_path} não encontrada!")
        return

    md_files = sorted(markdown_path.glob("*.md"))
    if not md_files:
        print(f"Nenhum arquivo Markdown (.md) encontrado em {markdown_path}")
        return
    
    criar_doc_com_conteudo(md_files)

# Execução
if __name__ == "__main__":
    main()